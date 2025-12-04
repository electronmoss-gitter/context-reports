"""
DOCX Parser - Extract text and metadata from Word earthing study reports
"""
from docx import Document
from pathlib import Path
from typing import Dict, List
import re

class DOCXParser:
    """Parse DOCX documents and extract structured content"""
    
    def __init__(self):
        self.section_patterns = {
            "executive_summary": r"(?i)(executive\s+summary|summary)",
            "site_description": r"(?i)(site\s+description|project\s+description|location)",
            "methodology": r"(?i)(methodology|method|approach|standards)",
            "soil_resistivity": r"(?i)(soil\s+resistivity|soil\s+test|wenner)",
            "earthing_design": r"(?i)(earthing\s+(system\s+)?design|grid\s+design|earth\s+electrode)",
            "calculations": r"(?i)(calculations?|analysis|design\s+calculations)",
            "touch_step": r"(?i)(touch\s+(and\s+)?step\s+potential|safety\s+analysis)",
            "compliance": r"(?i)(compliance|standards?\s+compliance|conformance)",
            "recommendations": r"(?i)(recommendations?|conclusions?)"
        }
    
    def parse(self, docx_path: str) -> Dict:
        """
        Parse a DOCX document and extract text with metadata
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dict with extracted text, sections, and metadata
        """
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")
        
        doc = Document(docx_path)
        
        # Extract all text
        full_text = self._extract_text(doc)
        
        # Extract metadata
        metadata = self._extract_metadata(doc, full_text, docx_path)
        
        # Split into sections
        sections = self._split_into_sections(doc)
        
        # Extract tables
        tables = self._extract_tables(doc)
        
        return {
            "full_text": full_text,
            "sections": sections,
            "metadata": metadata,
            "tables": tables,
            "source_file": str(docx_path)
        }
    
    def _extract_text(self, doc: Document) -> str:
        """Extract all text from document"""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return "\n\n".join(text_parts)
    
    def _extract_metadata(self, doc: Document, text: str, docx_path: Path) -> Dict:
        """Extract metadata from document properties and content"""
        metadata = {
            "filename": docx_path.name,
            "doc_type": "report"
        }
        
        # Extract from document properties
        core_props = doc.core_properties
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["created_date"] = str(core_props.created)
        if core_props.modified:
            metadata["modified_date"] = str(core_props.modified)
        if core_props.title:
            metadata["title"] = core_props.title
        
        # Extract project type
        if any(word in text.lower() for word in ["substation", "switchyard", "terminal"]):
            metadata["project_type"] = "substation"
        elif any(word in text.lower() for word in ["solar", "photovoltaic", "pv"]):
            metadata["project_type"] = "solar_farm"
        elif "wind" in text.lower():
            metadata["project_type"] = "wind_farm"
        else:
            metadata["project_type"] = "general"
        
        # Extract voltage level
        voltage_pattern = r"(\d+)\s*kV"
        voltage_matches = re.findall(voltage_pattern, text)
        if voltage_matches:
            voltages = [int(v) for v in voltage_matches]
            max_voltage = max(voltages)
            if max_voltage <= 1:
                metadata["voltage_level"] = "LV"
            elif max_voltage <= 66:
                metadata["voltage_level"] = "HV"
            else:
                metadata["voltage_level"] = "EHV"
        
        # Extract fault current range
        fault_pattern = r"(\d+(?:\.\d+)?)\s*kA"
        fault_matches = re.findall(fault_pattern, text)
        if fault_matches:
            fault_currents = [float(f) for f in fault_matches]
            max_fault = max(fault_currents)
            if max_fault < 10:
                metadata["fault_current_range"] = "0-10kA"
            elif max_fault < 50:
                metadata["fault_current_range"] = "10-50kA"
            else:
                metadata["fault_current_range"] = "50kA+"
        
        # Extract standards referenced
        standards = []
        if "AS/NZS 3000" in text or "AS/NZS3000" in text:
            standards.append("AS/NZS 3000")
        if "AS 2067" in text or "AS2067" in text:
            standards.append("AS 2067")
        if "IEEE 80" in text or "IEEE-80" in text:
            standards.append("IEEE 80")
        if "IEC 61936" in text:
            standards.append("IEC 61936")
        
        if standards:
            metadata["standards_referenced"] = standards
        
        return metadata
    
    def _split_into_sections(self, doc: Document) -> Dict[str, str]:
        """
        Split document into logical sections based on headings
        
        Returns:
            Dict mapping section_type to text content
        """
        sections = {}
        current_section = "introduction"
        current_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if paragraph is a heading
            is_heading = paragraph.style.name.startswith('Heading')
            
            # Check if text matches a section pattern
            section_found = False
            if is_heading or len(text) < 100:  # Headers are typically short
                for section_name, pattern in self.section_patterns.items():
                    if re.search(pattern, text):
                        # Save previous section
                        if current_text:
                            sections[current_section] = '\n'.join(current_text).strip()
                        
                        # Start new section
                        current_section = section_name
                        current_text = [text]
                        section_found = True
                        break
            
            if not section_found:
                current_text.append(text)
        
        # Save last section
        if current_text:
            sections[current_section] = '\n'.join(current_text).strip()
        
        return sections
    
    def _extract_tables(self, doc: Document) -> List[Dict]:
        """
        Extract tables from document
        
        Returns:
            List of tables as dicts
        """
        tables = []
        
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    "table_index": table_idx,
                    "data": table_data,
                    "row_count": len(table_data),
                    "col_count": len(table_data[0]) if table_data else 0
                })
        
        return tables


if __name__ == "__main__":
    # Test the parser
    parser = DOCXParser()
    
    # Example usage
    sample_path = Path("../data/historical_reports/sample_report.docx")
    if sample_path.exists():
        result = parser.parse(str(sample_path))
        print(f"Extracted {len(result['full_text'])} characters")
        print(f"Found {len(result['sections'])} sections")
        print(f"Found {len(result['tables'])} tables")
        print(f"Metadata: {result['metadata']}")